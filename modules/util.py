# 程序名：util.py
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

# ========== python 公共资源库 ==============
# 推迟类型注解解析，支持前向引用并减少运行时依赖
from __future__ import annotations
# 提供与解释器环境的交互能力（stderr、exit、sys.path 等）
import sys
# 提供日期时间。
import datetime
# 提供命令行参数解析能力
import argparse
# 提供面向对象的路径与文件操作
from pathlib import Path
# 导入 PDF 解析库（用于读取 PDF 并交给子模块处理）
import pdfplumber
# 用于读取 Excel 文件，并处理为 DataFrame 格式表格
import pandas as pd
from dataclasses import dataclass, field, asdict
# 提供类型注解所需的通用类型（列表、任意类型等）
from typing import List, Any, Iterator, Optional, Iterable, Union, Tuple
# 读取与解析 INI 配置文件
import configparser
# 操作系统级功能（路径、环境变量、文件与目录检测等）
import os
# PyMuPDF - 用于PDF文档操作（打开、解析、提取文本/图像等）
import fitz
# 导入 PaddleOCR 库的 PPStructureV3 类，用于版面分析和结构化文档识别（表格、段落等）
from paddleocr import PPStructureV3
# python-docx库的核心类，用于创建/修改Word文档（.docx格式）
from docx import Document 
from docx.text.paragraph import Paragraph
from docx.table import Table
# 给 docx.Document 起个别名，便于 isinstance 判断与类型标注
WordDocument = Document
from bs4 import BeautifulSoup
# 导入深拷贝工具
from copy import deepcopy
# 导入分页控制常量
from docx.enum.text import WD_BREAK

from docx.document import Document as WordDocument
# ========== End of  python 公共资源库 ==============

# ========== 软件项目环境目录 ==========
# 计算项目根目录（scripts/ 的上一级）
PROJECT_ROOT = Path(__file__).resolve().parents[1]
# 组装 src 目录路径
MODULES_DIR = PROJECT_ROOT / "modules"
# 将 modules 目录加入模块搜索路径（若尚未加入）
if str(MODULES_DIR) not in sys.path:
    sys.path.insert(0, str(MODULES_DIR))
# 将项目根目录加入模块搜索路径（若尚未加入）
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))
# ========== End of 软件项目环境目录 ==========

# ========== 配置类 ==========
@dataclass
class Config:
    # ========== 全局配置变量 ==========
    # 本处全局变量的含义详见doc-prep.ini文件对应变量的注释说明。
    # [Path]
    # 测试或投入运行，如果是测试，则输入文件访问test_path路径的文件，如果是运行，则访问input_path的文件。
    IS_TEST: bool = True
    # 输入文件路径
    INPUT_PATH: str = "data/"
    # 输出文件路径
    OUTPUT_PATH: str = "out/"
    # 测试文件目录
    TEST_PATH: str = "test/"
    # 日志文件路径
    LOG_FILE: str = "log/logfile.txt"
    # 批量处理文件列表名
    FILE_LIST_NAME: str = "file_list.txt"
    # 临时文件路径
    TEMP_FILE_PATH: str = "tmp/"
    # 中间过程 png 文件目录
    IMAGES_PATH: str  = "tmp/images/"

    # [File]
    # 是否拆分大文件。true 拆分；false 不拆分
    SPLIT_FILE: bool = False
    # 拆分文件阈值，大于阈值的文件即被拆分。拆分发生在pdf文件转换成word文件过程中。单位：MB，
    SPLIT_FILE_SIZE: int = 10
    # PDF文件转换word文件模式。可选 'structure_only（正文，表格）' 或 'smart_mixed（正文，表格，图片混排）'
    CONVERT_MODE: str = "structure_only"
    # 拆分文件合并模式：1 pdf文件合并（这种模式待验证）；2 word文件合并（目前这种模式有效）。
    MERGE_FILES_MODE: int = 2
    # pdf文件保存模式：neat 清洁干净保存模式（保存时间较长）；speed 快速保存模式（快速保存，但是文件内的垃圾较多，体积较大）
    SAVE_PDF_FILE_MODE: str = "speed"
    # 最终文件后缀名。用于补充最终输出文件的后缀名，例如：原文件ming_后缀.docx
    FILE_SUFFIX: str = "切片"

    # [Debug]
    # Debug 模式（0=静默，1=全部，2=仅警告）
    DEBUG_MODE: int = 0

    # [PdfCleanPolicy]
    # 删除文档封面
    CLEAN_COVER_PAGE: bool = True
    # 删除噪音页面（可能干扰AI有效切片的页面）：true 删除；false 不删除。
    CLEAN_NOISE_PAGES: bool = True
    # 删除 PDF 页眉关键词（如“前言,目录”）
    DELETE_PAGE_HEADER_NAMES: str = ""
    # 删除页眉页脚
    CLEAN_HEADER_FOOTER: bool = True
    # 页眉区域高度占页面总高的比例（前 10% 区域）
    TOP_RATIO: float = 0.1
    # 页脚区域从页面底部向上起的比例（后 10% 区域）
    BOTTOM_RATIO: float = 0.9
    # 判断页眉页脚重复行的前缀长度（单位：字符数）
    LINE_PREFIX_LEN: int = 20
    # 被认为是“共性前缀”的出现频率阈值（如 50% 页面中出现）
    FREQ_THRESHOLD: float = 0.5
    # 是否转换为 docx
    CONVERT_TO_DOCX: bool = True

    # [DocCleanPolicy]
    # Word 文档清洗策略
    FILE_CLEAN: bool = True
    # 清洗软回车“噪音”。
    CLEAN_SOFT_BREAK: bool = True
    # 清洗违例标识。
    CLEAN_VIOLATED_TAGS: bool = True
    # 违例定界符号，由于文档格式“噪音”，将本来应该另起一行的符号连接到上一行句子中了，
    # 例如：“步骤5 检查风扇模块是否存在异常。● 是 => 步骤6” 应该是：
    #       “步骤5 检查风扇模块是否存在异常。
    #         ● 是 => 步骤6”
    VIOLATED_TAGS: str = "●"
    # 清洗表格Table“噪音”
    CLEAN_TABLE: bool = True
    # 清洗表格后的输出文本表格格式。
    OUTPUT_TABLE_FORMAT: int = 1
    # Table标题列单元格内文字长度。
    TABLE_TITLE_MAX_LEN: int = 20
    # 识别文档标题的级别
    HEADING_DETECTION_LEVEL: int = 2
    # 文档内容标题最大长度
    HEADING_MAX_LEN: int = 60
    # 句子文本内容最小长度。小于这个长度的内容无效。
    TXT_MIN_LEN: int = 3
    # 清洗图片策略。
    CLEAN_PICS: bool = True
    # 合并被分隔的表格table：true 合并；false 不合并
    MERGE_TABLE: bool = True
    # 清洗空白段落：true 清洗；false 不清洗
    CLEAN_EMPTY_PARAGRAPH: bool = True

    #  [SliceTag]
    # 切片标记策略
    PARAGRAPH_MARK: int = 0
    SLICE_TAG: str = "/!!!!!"

    # 内部paddleocr v3 模型。
    PIPELINE: Any = None
    # 配置数据字典（doc_prep.ini）
    CONF_DICT: Any = None
    # ========= End of 全局配置变量 ===============

    # ========= 配置文件处理 ==========
    # 读取 INI 格式的配置文件，返回一个嵌套字典。
    def load_config_func(self, config_path: str) -> dict:
        """
        读取 INI 格式的配置文件，返回一个嵌套字典。
        参数：
            config_path: 配置文件路径
         返回：
            dict: 形如 {section: {key: value, ...}, ...} 的字典
        """
        config = configparser.ConfigParser()

        if not os.path.exists(config_path):
            print(f"❌ 配置文件不存在：{config_path}")
            return {}
        try:
            config.read(config_path, encoding="utf-8")
            print(f"✅ 成功加载配置文件：{config_path}")
        except Exception as e:
            print(f"❌ 配置文件加载失败：{e}")
            return {}
        # 转为嵌套字典
        config_dict = {section: dict(config.items(section)) for section in config.sections()}
        return config_dict

    #  将配置文件内容赋值给全局变量。
    def assign_config_to_globals(self, config: dict):
        """
        将配置文件内容赋值给全局变量
        """
        # [Path]        
        self.IS_TEST = config.get("Path", {}).get("is_test", self.IS_TEST)
        self.INPUT_PATH = config.get("Path", {}).get("input_path", self.INPUT_PATH)
        self.OUTPUT_PATH = config.get("Path", {}).get("output_path", self.OUTPUT_PATH)
        self.TEST_PATH = config.get("Path", {}).get("test_path", self.TEST_PATH)    
        self.FILE_LIST_NAME = config.get("Path", {}).get("file_list_name", self.FILE_LIST_NAME)
        self.LOG_FILE = config.get("Path", {}).get("log_file", self.LOG_FILE)
        self.TEMP_FILE_PATH = config.get("Path", {}).get("temp_file_path", self.TEMP_FILE_PATH)
        self.IMAGES_PATH = config.get("Path", {}).get("images_path", self.IMAGES_PATH)

        #[fILE]
        self.SPLIT_FILE = config.get("File", {}).get("split_file", "true").lower() == "true"
        self.SPLIT_FILE_SIZE = int(config.get("File", {}).get("split_file_size", self.SPLIT_FILE_SIZE))
        self.CONVERT_MODE = config.get("File", {}).get("convert_mode", self.CONVERT_MODE)
        self.MERGE_FILES_MODE = int(config.get("File", {}).get("merge_files_mode", self.MERGE_FILES_MODE))
        self.SAVE_PDF_FILE_MODE = config.get("File", {}).get("save_pdf_file_mode", self.SAVE_PDF_FILE_MODE)
        self.FILE_SUFFIX = config.get("File", {}).get("file_suffix", self.FILE_SUFFIX)

        # [Debug]
        self.DEBUG_MODE = int(config.get("Debug", {}).get("debug", self.DEBUG_MODE))

        # [PdfCleanPolicy]
        self.CLEAN_COVER_PAGE = config.get("PdfCleanPolicy", {}).get("clean_cover_pages", "true").lower() == "true"
        self.CLEAN_NOISE_PAGES = config.get("PdfCleanPolicy", {}).get("clean_noise_pages", "true").lower() == "true"
        self.DELETE_PAGE_HEADER_NAMES = config.get("PdfCleanPolicy", {}).get("delete_pages_header_name", self.DELETE_PAGE_HEADER_NAMES)
        self.CLEAN_HEADER_FOOTER = config.get("PdfCleanPolicy", {}).get("clean_page_header_footer", "true").lower() == "true"
        self.TOP_RATIO  = float(config.get("PdfCleanPolicy", {}).get("top_ratio", self.TOP_RATIO))
        self.BOTTOM_RATIO  = float(config.get("PdfCleanPolicy", {}).get("bottom_ratio", self.BOTTOM_RATIO))
        self.LINE_PREFIX_LEN  = int(config.get("PdfCleanPolicy", {}).get("line_prefix_len", self.LINE_PREFIX_LEN))
        self.FREQ_THRESHOLD  = float(config.get("PdfCleanPolicy", {}).get("freq_threshold", self.FREQ_THRESHOLD))
        self.CONVERT_TO_DOCX = config.get("PdfCleanPolicy", {}).get("if_convert_pdf_to_docx", "true").lower() == "true"

        # [DocCleanPolicy]
        self.FILE_CLEAN = config.get("DocCleanPolicy", {}).get("file_clean", "true").lower() == "true"
        self.CLEAN_SOFT_BREAK = config.get("DocCleanPolicy", {}).get("clean_soft_break_tag", "true").lower() == "true"
        self.CLEAN_VIOLATED_TAGS = config.get("DocCleanPolicy", {}).get("clean_violated_tags", "true").lower() == "true"
        self.VIOLATED_TAGS = config.get("DocCleanPolicy", {}).get("violated_tags", self.VIOLATED_TAGS)
        self.CLEAN_TABLE = config.get("DocCleanPolicy", {}).get("clean_table", "true").lower() == "true"
        self.OUTPUT_TABLE_FORMAT = int(config.get("DocCleanPolicy", {}).get("output_table_format", self.OUTPUT_TABLE_FORMAT))
        self.TABLE_TITLE_MAX_LEN = int(config.get("DocCleanPolicy", {}).get("table_title_max_len", self.TABLE_TITLE_MAX_LEN))
        self.HEADING_MAX_LEN = int(config.get("DocCleanPolicy", {}).get("heading_max_len", self.HEADING_MAX_LEN))
        self.HEADING_DETECTION_LEVEL = int(config.get("DocCleanPolicy", {}).get("heading_detection_level", self.HEADING_DETECTION_LEVEL))
        self.TXT_MIN_LEN = int(config.get("DocCleanPolicy", {}).get("txt_min_len", self.TXT_MIN_LEN))
        self.CLEAN_PICS = config.get("DocCleanPolicy", {}).get("clean_pics", "true").lower() == "true"
        self.MERGE_TABLE = config.get("DocCleanPolicy", {}).get("merge_table", "true").lower() == "true"
        self.CLEAN_EMPTY_PARAGRAPH = config.get("DocCleanPolicy", {}).get("clean_empty_paragraph", "true").lower() == "true"

        # [SliceTag]
        self.PARAGRAPH_MARK = int(config.get("SliceTag", {}).get("paragraph_mark", self.PARAGRAPH_MARK))
        self.SLICE_TAG = config.get("SliceTag", {}).get("slice_tag", self.SLICE_TAG)

    # 打印配置文件内容。
    def print_config_func(self, config: dict):
        """
        按格式打印嵌套字典配置内容
        """
        if not config:
            print("⚠️ 配置为空")
            return

        print("\n📋 打印显示配置文件内容：")
        for section, options in config.items():
            print(f"  [{section}]")
            for key, value in options.items():
                print(f"      {key} = {value}")
        print("✅ 配置文件内容打印完成")

    # 打印全局配置变量内容。
    def print_global_config_func(self):
        """
        打印当前所有全局配置变量的值，用于调试与确认赋值结果。
        """
        print("\n📋 打印显示当前全局配置变量：")

        print(f"  [Path]")        
        print(f"      IS_TEST = {self.IS_TEST}")        
        print(f"      INPUT_PATH = {self.INPUT_PATH}")
        print(f"      OUTPUT_PATH = {self.OUTPUT_PATH}")
        print(f"      TEST_PATH = {self.TEST_PATH}")        
        print(f"      FILE_LIST_NAME = {self.FILE_LIST_NAME}")
        print(f"      LOG_FILE = {self.LOG_FILE}")
        print(f"      TEMP_FILE_PATH = {self.TEMP_FILE_PATH}")
        print(f"      IMAGES_PATH = {self.IMAGES_PATH}")

        print("  [File]")
        print(f"      SPLIT_FILE = {self.SPLIT_FILE}")
        print(f"      SPLIT_FILE_SIZE = {self.SPLIT_FILE_SIZE} MB")
        print(f"      CONVERT_MODE = {self.CONVERT_MODE}")
        print(f"      MERGE_FILES_MODE = {self.MERGE_FILES_MODE}")
        print(f"      SAVE_PDF_FILE_MODE = {self.SAVE_PDF_FILE_MODE}")
        print(f"      FILE_SUFFIX = {self.FILE_SUFFIX}")

        print("  [Debug]")
        print(f"      DEBUG_MODE = {self.DEBUG_MODE}")

        print("  [PdfCleanPolicy]")
        print(f"      CLEAN_COVER_PAGE = {self.CLEAN_COVER_PAGE}")
        print(f"      CLEAN_NOISE_PAGES = {self.CLEAN_NOISE_PAGES}")
        print(f"      CLEAN_HEADER_FOOTER = {self.CLEAN_HEADER_FOOTER}")
        print(f"      DELETE_PAGE_HEADER_NAMES = {self.DELETE_PAGE_HEADER_NAMES}")
        print(f"      TOP_RATIO = {self.TOP_RATIO}")
        print(f"      BOTTOM_RATIO = {self.BOTTOM_RATIO}")
        print(f"      LINE_PREFIX_LEN = {self.LINE_PREFIX_LEN}")
        print(f"      FREQ_THRESHOLD = {self.FREQ_THRESHOLD}")
        print(f"      CONVERT_TO_DOCX = {self.CONVERT_TO_DOCX}")

        print("  [DocCleanPolicy]")
        print(f"      FILE_CLEAN = {self.FILE_CLEAN}")
        print(f"      CLEAN_SOFT_BREAK = {self.CLEAN_SOFT_BREAK}")
        print(f"      CLEAN_VIOLATED_TAGS = {self.CLEAN_VIOLATED_TAGS}")
        print(f"      VIOLATED_TAGS = {self.VIOLATED_TAGS}")
        print(f"      CLEAN_TABLE = {self.CLEAN_TABLE}")
        print(f"      OUTPUT_TABLE_FORMAT = {self.OUTPUT_TABLE_FORMAT}")
        print(f"      TABLE_TITLE_MAX_LEN = {self.TABLE_TITLE_MAX_LEN}")
        print(f"      TXT_MIN_LEN = {self.TXT_MIN_LEN}")
        print(f"      HEADING_MAX_LEN = {self.HEADING_MAX_LEN}")
        print(f"      HEADING_DETECTION_LEVEL = {self.HEADING_DETECTION_LEVEL}")
        print(f"      CLEAN_PICS = {self.CLEAN_PICS}")
        print(f"      MERGE_TABLE = {self.MERGE_TABLE}")
        print(f"      CLEAN_EMPTY_PARAGRAPH = {self.CLEAN_EMPTY_PARAGRAPH}")
        
        print("  [SliceTag]")
        print(f"      PARAGRAPH_MARK = {self.PARAGRAPH_MARK}")
        print(f"      SLICE_TAG = {self.SLICE_TAG}")

        print("  [非配置内部对象]")
        print(f"      PIPELINE = {self.PIPELINE}")
        print(f"      CONF_DICT = {self.CONF_DICT}")        
        print(f"✅ 全局变量打印完成")

# 初始化 cfg，提供一个初始化配置的通用函数。
def init_cfg_func(config_path: str):
    cfg = Config()
    # 加载配置文件。
    conf_dict = cfg.load_config_func(config_path) 
    cfg.assign_config_to_globals(conf_dict)
    # 创建paddleocr v3 模型。
    cfg.PIPELINE = False 
    cfg.CONF_DICT = conf_dict
    return cfg         
# ========== End of 配置类 ==========

# ========== class ParagBlockQ ==========
# 用来管理多个页面中的多个正文或表格段。
# 每个段落使用内部Element类进行封装。
# 本类提供添加、获取、打印等基础操作，支持 Word/Markdown 等文档结构的重建过程。
class ParagBlockQ:
    # 正文块元素。
    class Element:
        def __init__(self, page_index: int, block_index: int, block_type: str, content: str, bbox: List[int]):
            self.page_index = page_index     # 页码索引（从 0 开始）
            self.block_index = block_index    # 当前页中段落编号
            self.block_type = block_type     # 段落类型，如 TEXT、TITLE、TABLE
            self.content = content         # 段落内容文本
            self.bbox = bbox             # 段落的坐标框（如 [x0, y0, x1, y1]）
            
        def __repr__(self):
            return f"[{self.block_type}] 第{self.page_index+1}页-{self.block_index}段: {self.content[:30]}..."

    # 初始化。
    def __init__(self):
        self.blocks: List[ParagBlockQ.Element] = []  # 段落队列（顺序保存）

    # 向正文块队列中添加一个新的队列元素。
    def append(self, block: "ParagBlockQ.Element"):
        self.blocks.append(block)

    # 返回正文块总数。
    def __len__(self):
        return len(self.blocks)

    # 支持索引访问。
    def __getitem__(self, index: int) -> "ParagBlockQ.Element":
        return self.blocks[index]

    # 打印当前队列中所有段落块的内容（供调试与核查）。
    def print_all(self):
        print(f"✅ ===== 打印队列（ParagBlockQ）内容 =====")
        for i, block in enumerate(self.blocks):
            print(f"页码: {block.page_index}")
            print(f"段序: {block.block_index}")
            print(f"坐标: {block.bbox}")
            print(f"类型: {block.block_type}")
            print(f"内容: {block.content}")
            print(f"队列计数： {i}\n")
        print("  ===== 打印队列内容结束 =====")

    # 根据指定页码和字符串内容，在队列中查找匹配的段落索引（block_index），全词匹配
    def find_block_index_by_text(self, keyword: str, page_index: int) -> int:
        for block in self.blocks:
            if block.page_index == page_index and block.content.strip() == keyword.strip():
                return block.block_index  # 返回原始文档中的段落编号
        return -1  # 未找到匹配项

    # 根据页码获取该页的所有正文段落（返回列表）
    def get_blocks_by_page(self, page_index: int) -> List["ParagBlockQ.Element"]:
        return [block for block in self.blocks if block.page_index == page_index]

    # 删除指定页码和段落索引对应的正文块
    def remove_block(self, page_index: int, block_index: int) -> bool:
        for i, block in enumerate(self.blocks):
            if block.page_index == page_index and block.block_index == block_index:
                del self.blocks[i]
                return True
        return False
# ========== class ParagBlockQ结束 ==========

# ========== 列表类 ==========
@dataclass
class MyList:
    """通用列表类，支持任意类型元素，基于 Python 内置 list 封装"""
    def __init__(self):
        # 内部用 Python 的 list 存储数据
        self._data: List[Any] = []

    # === 增加元素 ===
    def append(self, item: Any) -> None:
        """在列表尾部添加元素"""
        self._data.append(item)

    def insert(self, index: int, item: Any) -> None:
        """在指定索引位置插入元素"""
        self._data.insert(index, item)

    # === 删除元素 ===
    def remove(self, item: Any) -> bool:
        """删除第一个匹配的元素，成功返回 True，否则返回 False"""
        try:
            self._data.remove(item)
            return True
        except ValueError:
            return False

    def remove_at(self, index: int) -> None:
        """删除指定索引位置的元素"""
        if index < 0 or index >= len(self._data):
            raise IndexError("索引超出范围")
        del self._data[index]

    # === 获取元素 ===
    def __getitem__(self, index: int) -> Any:
        return self._data[index]

    def __setitem__(self, index: int, value: Any) -> None:
        self._data[index] = value

    # === 遍历与长度 ===
    def __iter__(self) -> Iterator[Any]:
        return iter(self._data)

    def __len__(self) -> int:
        return len(self._data)

    # === 其他功能 ===
    def clear(self) -> None:
        """清空列表"""
        self._data.clear()

    def to_list(self) -> List[Any]:
        """返回 Python 内置 list"""
        return list(self._data)

    def __repr__(self) -> str:
        return f"MyList({self._data})"

# 应用示例。
"""
cl = MyList()

# 添加不同类型的元素
cl.append("hello")
cl.append([1, 2, 3])
cl.append({"a": 10})
cl.append(42)

print(cl)  # MyList(['hello', [1, 2, 3], {'a': 10}, 42])

# 插入元素
cl.insert(2, "插入位置2")
print(cl)  # MyList(['hello', [1, 2, 3], '插入位置2', {'a': 10}, 42])

# 删除元素
cl.remove("hello")
print(cl)  # MyList([[1, 2, 3], '插入位置2', {'a': 10}, 42])

# 删除指定索引
cl.remove_at(1)
print(cl)  # MyList([[1, 2, 3], {'a': 10}, 42])

# 遍历
for item in cl:
    print("元素:", item)

# 获取索引
print("索引2元素:", cl[2])  # 42
"""
# ========== End of 列表类 ==========

# ========== 导入 PaddleOCR v3.1.0 的OCR识别模型 ==========
# 函数用途：
# 创建“安全默认”的 PPStructureV3（关闭公式/图表等大模型，避免 OOM）
# —— 替换脚本中的 create_safe_ppstructure_v3_func() —— 
# 安全创建 PPStructureV3 的函数，确保 PaddleOCR 初始化时不会因未知参数报错
def create_safe_ppstructure_v3_func():
    """
    安全创建 PPStructureV3。自动过滤 PaddleOCR 不支持的关键字参数，避免
    ValueError: Unknown argument: xxx
    """
    # === 原形参改为函数内部赋值 ===
    # 是否启用版面区域检测
    use_region_detection: bool = True
    # 是否启用表格识别
    use_table_recognition: bool = True
    # 是否启用公式识别
    use_formula_recognition: bool = False
    # 是否启用图表识别
    use_chart_recognition: bool = False
    # 是否启用印章识别
    use_seal_recognition: bool = False
    # 额外参数（上层可扩展，但这里默认置空）
    kwargs = {}
    print(f"\n✅ 正在导入百度飞桨 OCR PPStructureV3 模型......")
    # 尝试导入 PPStructureV3 模型
    try:
        from paddleocr import PPStructureV3
    except Exception as e:
        # 如果导入失败，抛出运行时错误
        raise RuntimeError(f"导入 PaddleOCR 失败：{e}")
    # 定义支持的关键字参数集合（白名单）
    supported_keys = {
        "use_region_detection",
        "use_table_recognition",
        "use_formula_recognition",
        "use_chart_recognition",
        "use_seal_recognition",
        # 如环境支持更多开关，可在这里补充
    }
    # 构建初始参数字典，填入基础配置
    kw = {
        "use_region_detection": use_region_detection,
        "use_table_recognition": use_table_recognition,
        "use_formula_recognition": use_formula_recognition,
        "use_chart_recognition": use_chart_recognition,
        "use_seal_recognition": use_seal_recognition,
    }
    # 遍历额外传入的参数 kwargs
    for k, v in kwargs.items():
        # 如果参数在支持的集合内，就覆盖到 kw 中
        if k in supported_keys:
            kw[k] = v
        else:
            # 否则打印警告，不中断程序
            print(f"⚠️  PPStructureV3.__init__ 不支持参数: {k}（已忽略）")
    # 打印最终传入构造器的参数，方便调试
    print(f"最终传入构造器的参数: {kw}")
    # 尝试用过滤后的参数字典创建 PPStructureV3
    try:
        pipeline = PPStructureV3(**kw)
    except TypeError as e:
        # 如果参数仍然不兼容，则提示并使用最简参数重试
        print(f"⚠️  PPStructureV3 参数不兼容，改用极简构造重试：{e}")
        pipeline = PPStructureV3(
            use_region_detection=use_region_detection,
            use_table_recognition=use_table_recognition,
            use_formula_recognition=use_formula_recognition,
            use_chart_recognition=use_chart_recognition,
            use_seal_recognition=use_seal_recognition,
        )
    # 返回创建好的 PPStructureV3 实例
    print(f"✅ 百度飞桨 OCR PPStructureV3 模型导入完毕")
    return pipeline

# 函数用途：
#   创建一个带安全默认参数的 PPStructureV3 OCR 管线
# 特性：
#   - 自动过滤当前版本不支持的参数，避免 Unknown argument 报错
#   - 默认关闭公式/图表/印章识别，降低内存占用
#   - 默认限制检测图片的最小边长，避免高分辨率导致 OOM
def 备用_create_safe_ppstructure_v3_func():
    # 配置参数。未来这里可以改成从配置文件或全局变量读取
    use_table_recognition = True
    use_region_detection = True
    use_formula_recognition = False
    use_chart_recognition = False
    use_seal_recognition = False
    limit_side_len = 1600
    print("\n✅ 初始化PaddleOCR3.1.0 模型 创建安全 PPStructureV3 OCR 管线")
    # 构造理想参数字典（有些版本可能不支持部分参数）
    desired_kwargs = {
        "use_region_detection": use_region_detection,
        "use_table_recognition": use_table_recognition,
        "use_formula_recognition": use_formula_recognition,
        "use_chart_recognition": use_chart_recognition,
        "use_seal_recognition": use_seal_recognition,
        "text_det_params": {
            "limit_side_len": int(limit_side_len),
            "limit_type": "min",
        },
    }
    # 获取当前版本 PPStructureV3.__init__ 支持的参数名
    init_params = set(signature(PPStructureV3.__init__).parameters.keys())
    # 只保留支持的参数
    safe_kwargs = {}
    for k, v in desired_kwargs.items():
        if k in init_params:
            safe_kwargs[k] = v
        else:
            print(f"[警告] PPStructureV3.__init__ 不支持参数: {k}（已忽略）")
    # 打印最终传入的参数（便于调试）
    print(f"最终传入构造器的参数: {safe_kwargs}")
    # 创建管线实例
    pipeline = PPStructureV3(**safe_kwargs)
    print("✅ 初始化PaddleOCR3.1.0 模型完成，管线实例创建完成")   
    # 返回管线对象
    return pipeline
# ========== End of 导入 PaddleOCR v3.1.0 的OCR识别模型 ==========

# ========== 文件杂项 ==========
# 根据输入的文件路径，生成目标文件路径。
def gen_target_file_name_func(input_file_path: str, target_dir: str, suffix: str) -> str:
    """
    根据输入文件路径和后缀名，生成“原文件名 + -后缀”的输出文件路径。
    参数：
        input_file_path: str，原始文件路径，例如：/home/ubuntu/slice/proj/data/测试.docx
        target_dir: str, 目标目录，例如： /home/ubuntu/slice/proj/out
        suffix: str，要添加的自定义后缀，
                如果带扩展名（如 "修改.pdf"），则覆盖原始扩展名；
                如果不带扩展名（如 "修改"），则保留原始扩展名。
    返回：
        str，生成的新路径，例如：
            /home/ubuntu/slice/proj/out/测试_修改.docx
            /home/ubuntu/slice/proj/out/测试_修改.pdf
    """
    # 分离路径、文件名和扩展名
    print(f"\n✅ 参考原文件路径: {input_file_path}")
    dir_name, base_name = os.path.split(input_file_path)
    name, ext = os.path.splitext(base_name)
    # 判断 suffix 是否自带扩展名
    suffix_name, suffix_ext = os.path.splitext(suffix)
    # 如果 suffix 自带扩展名，则替换原始扩展名
    if suffix_ext:
        new_name = f"{name}_{suffix_name}{suffix_ext}"
    # 否则，保留原始扩展名
    else:
        new_name = f"{name}_{suffix}{ext}"
    # 拼接输出文件路径
    output_file_path = os.path.join(target_dir, new_name)
    print(f"✅ 生成改名后的文件路径: {output_file_path}")
    # 返回完整路径
    return output_file_path

# 根据输入文件路径，生成“原文件名 + 后缀”的输出路径，支持修改扩展名
def gen_output_file_name_func(input_file_name: str, out_path: str, suffix: str, new_ext: str = "") -> str:
    """
    根据输入文件路径，生成“原文件名 + 后缀”的输出路径。
    参数：
        input_file_name: str，原始文件路径，如 /path/to/测试.docx
        out_path: str，保存目录，如 tmp/
        suffix: str，追加的文件名后缀，如 "中间文件"
        new_ext: str，可选的新扩展名（如 ".pdf"、".docx"），默认为空，表示使用原始扩展名。
    返回：
        str，生成的新文件完整路径，如 /tmp/测试_中间文件.docx
    """
    # 拆分路径与扩展名
    dir_name, base_name = os.path.split(input_file_name)
    name, ext = os.path.splitext(base_name)
    # 使用指定的保存路径
    dir_name = out_path
    # 使用新扩展名（如果有）
    final_ext = new_ext if new_ext else ext
    # 拼接新文件名
    if suffix == "":
        new_name = f"{name}{final_ext}"
    else:
        new_name = f"{name}_{suffix}{final_ext}"
    print(f"\n✅ 生成新文件名：{new_name}")
    # 拼接完整路径
    return os.path.join(dir_name, new_name)

# 遍历指定目录，返回所有 .png 文件的完整路径列表（按文件名升序排序）
# 参数:
# directory (str): 要遍历的根目录路径
# recursive (bool): 是否递归遍历子目录，默认 False
#返回:
# List[str]: 所有找到的 PNG 文件的路径列表
def get_png_files_func(directory: str, recursive: bool = False) -> List[str]:
    # 初始化 PNG 文件路径列表
    png_files = []
    # 如果目录不存在，直接返回空列表
    if not os.path.isdir(directory):
        print(f"⚠️ 目录不存在：{directory}（返回空列表）")
        return png_files
    if recursive:
        # ✅ 递归遍历目录和子目录
        for root, _, files in os.walk(directory):
            for file in files:
                # 判断文件扩展名是否为 .png（忽略大小写）
                if file.lower().endswith('.png'):
                    # 拼接文件完整路径
                    file_path = os.path.join(root, file)
                    # 加入文件列表
                    png_files.append(file_path)
    else:
        # ✅ 只遍历当前目录
        for file in os.listdir(directory):
            # 判断文件扩展名是否为 .png
            if file.lower().endswith('.png'):
                # 拼接文件完整路径
                file_path = os.path.join(directory, file)
                # 加入文件列表
                png_files.append(file_path)
    # ✅ 按文件名升序排序（不考虑路径，只看文件名部分）
    png_files.sort(key=lambda x: os.path.basename(x))
    # 打印显示列表文件，
    print_png_files_func(png_files)
    # 打印统计结果
    print(f"✅ 在目录 {directory} 中找到 {len(png_files)} 个 PNG 文件。 (recursive={recursive})")
    return png_files

def 备份_get_png_files_func(directory: str, recursive: bool = False) -> List[str]:
    """
    遍历指定目录，返回所有 .png 文件的完整路径列表
    参数:
        directory (str): 要遍历的根目录路径
        recursive (bool): 是否递归遍历子目录，默认 True

    返回:
        List[str]: 所有找到的 PNG 文件的路径列表
    """
    png_files = []
    if not os.path.isdir(directory):
        print(f"⚠️ 目录不存在：{directory}（返回空列表）")
        return png_files
    if recursive:
        # 递归遍历目录和子目录
        for root, _, files in os.walk(directory):
            for file in files:
                if file.lower().endswith('.png'):
                    file_path = os.path.join(root, file)
                    png_files.append(file_path)
    else:
        # 只遍历当前目录
        for file in os.listdir(directory):
            if file.lower().endswith('.png'):
                file_path = os.path.join(directory, file)
                png_files.append(file_path)
    print(f"✅ 在目录 {directory} 中找到 {len(png_files)} 个 PNG 文件。 (recursive={recursive})")
    return png_files

# 加载pdf文件。
def load_pdf_file_func(pdf_path):
    """
    加载 PDF 文档
    """
    try:
        doc = fitz.open(pdf_path)
        print(f"✅  成功加载 PDF 文件：{pdf_path}")
        return doc
    except Exception as e:
        print(f"❌  加载 PDF 文件失败：{e}")
        return None

# 加载 Word 文档（.docx 格式）
def load_docx_file_func(doc_path: str) -> Document:
    try:
        # 尝试加载 Word 文档
        doc = Document(doc_path)
        # 打印加载成功信息
        print(f"✅ 成功加载文档：{doc_path}")
        return doc
    except Exception as e:
        # 打印加载失败信息
        print(f"❌ 无法加载文档：{doc_path}，错误信息：{e}")
        return None

# 加载 excel 文档（.docx 格式）
def load_excel_file_func(file_path: str) -> pd.ExcelFile:
    """
    读取 Excel 文件为 ExcelFile 对象。
    :param file_path: Excel 文件路径
    :return: pandas.ExcelFile 对象（若文件不存在则抛出异常）
    """
    if not os.path.exists(file_path):
        print(f"❌ 文件未找到：{file_path}")
        return None
    try:
        excel_file = pd.ExcelFile(file_path)
        print(f"✅ 成功读取 Excel 文件：{file_path}")
        return excel_file
    except Exception as e:
        print(f"❌ 读取 Excel 文件出错：{e}")
        return None

# 克隆文档对象，生成一个新的副本
def clone_doc_func(doc: Document) -> Document:
    # 创建一个临时文件，后缀为 .docx
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        # 获取临时文件路径
        temp_path = tmp.name
        # 将源文档保存到临时文件
        doc.save(temp_path)
    # 从临时文件加载为新文档对象
    new_doc = Document(temp_path)
    # 删除临时文件，避免文件残留
    os.remove(temp_path)
    return new_doc

# 删除目录中的全部文件。
def remove_path_files_func(target_dir: str):
    print(f"\n✅ 删除{target_dir}目录下的全部文件")
    # 遍历目录中的所有文件
    for file_name in os.listdir(target_dir):
        # 拼接完整路径
        file_path = os.path.join(target_dir, file_name)
        # 如果是文件就删除
        if os.path.isfile(file_path):
            os.remove(file_path)    

# 删除当前目录下（包含子目录下）的全部文件。
def remove_path_recursio_files_func(target_dir: str):
    import os
    print(f"\n✅ 递归删除 {target_dir} 目录及其子目录下的全部文件")

    for root, dirs, files in os.walk(target_dir):
        for file_name in files:
            file_path = os.path.join(root, file_name)
            try:
                os.remove(file_path)
                print(f"🗑️ 已删除文件：{file_path}")
            except Exception as e:
                print(f"⚠️ 删除失败：{file_path}，原因：{e}")

# 打印 PNG 文件列表，显示索引和文件名（用于检查排序结果）
def print_png_files_func(png_files: List[str]):
    if not png_files:
        print("⚠️ PNG 文件列表为空。")
        return
    # 打印列表头部说明
    print("\n📂 打印 PNG 文件列表（已按文件名升序排序）：")
    # 遍历文件列表，显示序号、文件名和完整路径
    for idx, path in enumerate(png_files, start=1):
        # idx:03d → 序号占 3 位，不足补 0
        print(f"{idx:03d}. {os.path.basename(path)}  ({path})")
    print("   ========== 打印完毕 ==========")

# 打印 Word 文档对象的结构化内容
# 参数：
#   doc : docx.Document 对象
# 功能：
#   按顺序打印所有段落和表格内容，用于调试或结构验证
def print_docx_func(doc: Document):
    print(f"\n✅ 打印 doc 文档对象的内容")
    # 初始化段落和表格的序号计数器
    para_count = 0
    table_count = 0
    # 遍历 Word 文档的底层 block（段落或表格）
    for block in doc.element.body:
        # 如果是段落节点（<w:p>）
        if block.tag.endswith("}p"):
            # 转换为 docx 的 Paragraph 对象
            para = Paragraph(block, doc)
            # 增加段落编号
            para_count += 1
            # 打印段落内容，带编号
            print(f"[段落 {para_count}] {para.text}")
        # 如果是表格节点（<w:tbl>）
        elif block.tag.endswith("}tbl"):
            # 转换为 docx 的 Table 对象
            table = Table(block, doc)
            # 增加表格编号
            table_count += 1
            # 打印表格标记
            print(f"[表格 {table_count}]")
            # 遍历表格的每一行
            for row in table.rows:
                # 提取每个单元格内容，并拼接成一行文本
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                print(f"    {row_text}")
    print(f"\n📋 总计：{para_count} 个正文段落，{table_count} 个表格")
    print(f"✅ 打印 doc 文档对象内容结束")

# 检查 word 文档对象 Document 是否有效
def check_docx_func(doc: Document):
    if doc is None:
        print(f"\n❌ 错误：传入的 doc 参数为 None，请确认文档是否正确加载")
    else:
        print(f"\n✅ 传入的 doc 参数有效，文档已正确加载")
    input(f"暂停 .......")

# 保存 pdf文档或 word 文档到指定路径。能够识别处理保存pdf文档与word文档。
#   参数：
#        doc: Word 或 PDF 文档对象（docx.Document 或 fitz.Document）
#        file_path: 保存路径
def save_doc_func(doc: Union[WordDocument, fitz.Document], file_path: str) -> None:
    """
    保存 PDF 或 Word 文档到指定路径。
    参数：
        doc: Word 或 PDF 文档对象（docx.Document 或 fitz.Document）
        file_path: 保存路径
    """
    try:
        # 判断是否为 Word 文档
        if isinstance(doc, WordDocument):
            doc.save(file_path)
            print(f"✅ Word文档已保存：{file_path}")
        # 判断是否为 PDF 文档
        elif isinstance(doc, type(fitz.open())):
            if Config.SAVE_PDF_FILE_MODE == "speed":
                print(f"⏳ 正在以快速方式(speed模式)保存pdf文件，请等待......")
                doc.save(file_path)
                doc.close()
                print(f"✅ PDF文档已保存：{file_path}")
            elif Config.SAVE_PDF_FILE_MODE == "neat":
                print(f"⏳ 正在以清洁干净方式(neat模式)保存pdf文件，需要耐心等待较长时间......")
                # 说明：在测试时发现，pdf文档在经过清洗（删除页眉页脚，删除前言，目录章节后，文件占用空间体积会明显变大，甚至增大至6、7倍。
                # 询问ChatGPT，告知会保存很多无用垃圾。而使用doc.save(file_path)指令正是产生这种现象的原因。需要使用更好的指令。2025-07-22
                # 清洗保存pdf文档。garbage=4 清除的是 无引用的“垃圾对象”；deflate=True 是一种无损压缩算法；
                # clean=True 是对 PDF 的结构重构，而非内容变更。
                doc.save(file_path, garbage=4, deflate=True, clean=True)
                doc.close()
                print(f"✅ PDF文档已保存：{file_path}")
            else:
                print(f"⚠️  pdf文档保存模式错误：SAVE_PDF_FILE_MODE = {SAVE_PDF_FILE_MODE}，无法保存。")
        else:
            print(f"⚠️ 未知文档类型，无法保存。")
    except Exception as e:
        print(f"❌ 保存失败：{e}")

# 遍历子 Document 实例，合并成一个最终文档
#   参数:
#        doc_list: list[Document]  子文档列表
#   返回:
#        Document 合并后的最终文档
# ========== 合并子文档对象成最终文档 ==========
# 自定义合并函数：保留表格结构，避免生成多余空白页
# 说明：
# 在 doc_prep 包缺失或其 util 模块未定义 merge_documents_func 时，
# 本函数可作为替代。它在合并多个子文档时，去除初始文档默认空段落，
# 跳过完全空白的子文档，并只在相邻两个非空文档之间插入分页符，
# 以防止合并后的文档开头出现连续空白页。同时使用深拷贝追加底层元素，
# 保留表格中单元格合并等结构信息。
def merge_documents_func(doc_list: list[Document]) -> Document:
    """
    将多个子文档合并为一个文档，同时保留表格格式并避免开头产生多余空白页。

    参数：
        doc_list: List[Document] 子文档列表。

    返回：
        Document 合并后的文档对象。
    """
    # 创建最终文档对象
    final_doc = Document()
    # 清空文档默认的空段落，避免第一页面出现空白
    final_doc._element.body.clear_content()
    # 初始化一个列表用于存储非空文档的索引
    non_empty_indices: list[int] = []
    # 预先扫描 doc_list，找出包含有效内容的文档索引
    for idx, sub_doc in enumerate(doc_list):
        # 标记是否有内容
        has_content = False
        # 检查段落中是否存在非空文本
        for para in sub_doc.paragraphs:
            if para.text.strip():
                has_content = True
                break
        # 如果没有非空文本但含有表格，也算有内容
        if (not has_content) and sub_doc.tables:
            has_content = True
        # 若文档确实包含内容，则记录其索引
        if has_content:
            non_empty_indices.append(idx)
    # 如果没有任何非空文档，则直接返回空白 final_doc
    if not non_empty_indices:
        return final_doc
    # 遍历子文档列表以构建合并内容
    for idx, sub_doc in enumerate(doc_list):
        # 如果此文档无内容则跳过
        if idx not in non_empty_indices:
            continue
        # 遍历文档中块级元素（段落、表格）
        for child in sub_doc._element.body.iterchildren():
            # 跳过节属性节点以防止页面设置冲突
            if child.tag.endswith('sectPr'):
                continue
            # 将元素深拷贝后追加到 final_doc
            final_doc._element.body.append(deepcopy(child))
    # 返回合并完成的文档
    return final_doc
# ========== End of 文件杂项 ==========

# ========= 判断输入文件类型 ==========
# 判断指定文件是否是一个合法的 PDF 文件，并打印检查过程。
def is_pdf_file_func(file_path: str) -> bool:
    """
    判断指定文件是否是一个合法的 PDF 文件，并打印检查过程。
    参数：
        file_path: 文件路径
    返回：
        True：是 PDF 且能成功打开；
        False：不是 PDF 或打开失败
    """
    print(f"⚠️  正在检查文件是否为有效 PDF 文件 ...")
    path = Path(file_path)

    # 检查文件是否存在
    if not path.is_file():
        print(f"❌ 文件不存在：{file_path}")
        return False
    else:
        print(f"✅ 检查文件存在：{file_path}")

    # 检查文件扩展名
    if path.suffix.lower() != ".pdf":
        print(f"❌ 文件扩展名不是 .pdf（实际为 {path.suffix}）")
        return False
    else:
        print("✅ 文件扩展名为 .pdf")

    # 尝试使用 PyMuPDF 打开文件
    try:
        doc = fitz.open(file_path)
        if doc.is_pdf:
            print(f"✅ 文件成功打开，确认是 PDF 格式")
            doc.close()
            return True
        else:
            print(f"❌ 文件打开成功，但不是 PDF 格式")
            doc.close()
            return False
    except Exception as e:
        print(f"❌ 打开文件失败：{e}")
        return False

# 判断指定文件是否是一个合法的 Word (.docx) 文件
def is_word_file_func(file_path: str) -> bool:
    print(f"⚠️  正在检查文件是否为有效 Word 文件 ...")
    path = Path(file_path)
    if not path.is_file():
        print(f"❌ 文件不存在：{file_path}")
        return False
    else:
        print(f"✅ 检查文件存在：{file_path}")
    if path.suffix.lower() != ".docx":
        print(f"❌ 文件扩展名不是 .docx（实际为 {path.suffix}）")
        return False
    else:
        print("✅ 文件扩展名为 .docx")
    try:
        doc = Document(file_path)
        _ = doc.paragraphs  # 尝试访问段落，确认结构正常
        print(f"✅ 文件成功打开，确认是 Word 格式")
        return True
    except Exception as e:
        print(f"❌ 打开 Word 文件失败：{e}")
        return False
# ========= End of 判断输入文件类型 ==========

# ========== 检查 word 文档内的表格结构 ==========
# 定义一个内部工具函数，用于统计表格内的 gridSpan 和 vMerge 标签数量
def _count_spans(tbl_element) -> Tuple[int, int]:
    # 初始化 gridSpan 与 vMerge 计数器
    gridspan = 0
    vmerge = 0
    # 遍历表格元素的所有子元素
    for el in tbl_element.iter():
        # 获取当前元素的标签
        tag = el.tag
        # 确保标签是字符串类型
        if isinstance(tag, str):
            # 如果标签以 gridSpan 结尾，说明是列合并，计数加一
            if tag.endswith('gridSpan'):
                gridspan += 1
            # 如果标签以 vMerge 结尾，说明是行合并，计数加一
            elif tag.endswith('vMerge'):
                vmerge += 1
    # 返回 gridSpan 和 vMerge 的数量
    return gridspan, vmerge

# 定义主函数：检查一个或多个 DOCX 文件的表格结构
def inspect_docx_tables(
    docx_paths: Union[str, Iterable[str]],
    save_csv: Optional[str] = None,
    print_details: bool = True,
):
    # 如果输入的是单个字符串路径，则转换为列表
    if isinstance(docx_paths, (str, bytes, os.PathLike)):
        paths: List[str] = [str(docx_paths)]
    # 否则将其转换为字符串列表
    else:
        paths = [str(p) for p in docx_paths]
    # 初始化结果列表
    results: List[dict] = []
    # 遍历所有待检查的 DOCX 文件路径
    for path in paths:
        # 如果文件不存在则跳过
        if not os.path.isfile(path):
            if print_details:
                print(f"⚠️  跳过不存在的文件: {path}")
            continue
        # 尝试打开 DOCX 文件
        try:
            doc = Document(path)
        except Exception as e:
            if print_details:
                print(f"⚠️  无法打开 DOCX: {path} -> {e}")
            continue
        # 如果文档没有表格，则添加一行默认结果
        if not doc.tables:
            row = {
                "file": os.path.basename(path),
                "table_index": None,
                "rows": 0,
                "cols": 0,
                "gridSpan_count": 0,
                "vMerge_count": 0,
            }
            results.append(row)
            if print_details:
                print(f"\n===== TABLE INSPECT: {os.path.basename(path)} =====")
                print("tables: 0, paragraphs:", len(doc.paragraphs))
                print("  (无表格)")
            continue
        # 如果文档包含表格，打印文档级别统计信息
        if print_details:
            print(f"\n===== TABLE INSPECT: {os.path.basename(path)} =====")
            print(f"tables: {len(doc.tables)}, paragraphs: {len(doc.paragraphs)}")
        # 遍历文档中的所有表格
        for ti, tbl in enumerate(doc.tables):
            # 调用内部函数统计 gridSpan 和 vMerge 数量
            gs, vm = _count_spans(tbl._element)
            # 生成一行统计结果
            row = {
                "file": os.path.basename(path),
                "table_index": ti,
                "rows": len(tbl.rows),
                "cols": len(tbl.columns),
                "gridSpan_count": gs,
                "vMerge_count": vm,
            }
            # 将结果追加到列表中
            results.append(row)
            # 打印该表格的统计结果
            if print_details:
                print(f"  - table#{ti}: rows={row['rows']}, cols={row['cols']}, gridSpan={gs}, vMerge={vm}")
    # 如果用户要求保存为 CSV 文件
    if save_csv:
        try:
            # 尝试使用 pandas 保存
            import pandas as pd
            df = pd.DataFrame(results)
            os.makedirs(os.path.dirname(save_csv) or ".", exist_ok=True)
            df.to_csv(save_csv, index=False, encoding="utf-8-sig")
            if print_details:
                print(f"\n💾 已保存体检明细 CSV：{save_csv}")
        except Exception:
            # 如果 pandas 不可用，退回标准库 csv
            import csv
            os.makedirs(os.path.dirname(save_csv) or ".", exist_ok=True)
            with open(save_csv, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.DictWriter(
                    f,
                    fieldnames=["file", "table_index", "rows", "cols", "gridSpan_count", "vMerge_count"]
                )
                writer.writeheader()
                for r in results:
                    writer.writerow(r)
            if print_details:
                print(f"\n💾 已保存体检明细 CSV（使用 csv 标准库）：{save_csv}")
    # 最后尝试返回 pandas DataFrame，若失败则返回原始列表
    try:
        import pandas as pd
        return pd.DataFrame(results)
    except Exception:
        return results

# 用法示例
"""
# 1) 单文件
inspect_docx_tables("TaiShan200安装指南-3栏_p02_c02_中间.docx")

# 2) 多文件 + 保存 CSV
files = [
    "TaiShan200安装指南-3栏_p01_c01_中间.docx",
    "TaiShan200安装指南-3栏_p01_c02_中间.docx",
    "TaiShan200安装指南-3栏_p01_c03_中间.docx",
    "TaiShan200安装指南-3栏_p02_c01_中间.docx",
    "TaiShan200安装指南-3栏_p02_c02_中间.docx",
]
df = inspect_docx_tables(files, save_csv="./_debug/table_inspect.csv")
print(df)
"""
# ========== End of 检查 word 文档内的表格结构 ==========

# ========== 日志 Logger 类 ==========
"""
功能描述：
---------------------------------------
1. 提供 Logger 类，用于日志打印与保存；
2. 控制台输出带颜色、时间戳；
3. 自动创建 logs/ 目录并保存日志文件；
4. 供整个项目的各模块调用。
---------------------------------------
"""
class Logger:
    """
    日志工具类
    -------------------------
    支持 info / warn / error 三种日志级别
    """

    def __init__(self, log_dir: str = "../logs"):
        # 日志保存目录
        self.log_dir = os.path.abspath(log_dir)
        os.makedirs(self.log_dir, exist_ok=True)

        # 日志文件路径（按日期命名）
        date_str = datetime.datetime.now().strftime("%Y-%m-%d")
        self.log_file = os.path.join(self.log_dir, f"run_{date_str}.log")

        # 打印初始化信息
        self._write_to_console("Logger", "日志系统初始化成功", level="INFO")

    # ---------------------------
    # 控制台颜色定义
    # ---------------------------
    COLORS = {
        "INFO": "\033[92m",   # 绿色
        "WARN": "\033[94m",   # 蓝色
        "ERROR": "\033[91m",  # 红色
        "RESET": "\033[0m",   # 颜色重置
    }

    # ---------------------------
    # 内部方法：格式化消息
    # ---------------------------
    def _format_message(self, level: str, message: str, log_tag: str = None) -> str:
        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        prefix = f"[{now}] [{level}]"
        if log_tag:
            prefix += f" [{log_tag}]"
        return f"{prefix} {message}"

    # ---------------------------
    # 内部方法：输出到控制台
    # ---------------------------
    def _write_to_console(self, log_tag: str, message: str, level: str = "INFO"):
        color = self.COLORS.get(level, "")
        reset = self.COLORS["RESET"]
        formatted = self._format_message(level, message, log_tag)
        print(f"{color}{formatted}{reset}")

    # ---------------------------
    # 内部方法：写入日志文件
    # ---------------------------
    def _write_to_file(self, message: str):
        with open(self.log_file, "a", encoding="utf-8") as f:
            f.write(message + "\n")

    # ---------------------------
    # 公共方法：INFO
    # ---------------------------
    def info(self, message: str, log_tag: str = None):
        formatted = self._format_message("INFO", message, log_tag)
        self._write_to_console(log_tag, message, level="INFO")
        self._write_to_file(formatted)

    # ---------------------------
    # 公共方法：WARN
    # ---------------------------
    def warn(self, message: str, log_tag: str = None):
        formatted = self._format_message("WARN", message, log_tag)
        self._write_to_console(log_tag, message, level="WARN")
        self._write_to_file(formatted)

    # ---------------------------
    # 公共方法：ERROR
    # ---------------------------
    def error(self, message: str, log_tag: str = None):
        formatted = self._format_message("ERROR", message, log_tag)
        self._write_to_console(log_tag, message, level="ERROR")
        self._write_to_file(formatted)

    # ---------------------------
    # 公共方法：show_config
    # ---------------------------
    def show_config(self, config: configparser.ConfigParser, log_tag: str = "Config") -> None:
        """
        显示 configparser 配置内容，用于调试。

        参数:
            config: configparser.ConfigParser 实例
            log_tag: 日志中显示的模块标签
        """
        if not config.sections():
            self.warning("⚠️ 配置为空或读取失败", log_tag=log_tag)
            return

        for section in config.sections():
            self.info(f"[{section}]", log_tag=log_tag)
            for key, value in config.items(section):
                self.info(f"{key} = {value}", log_tag=log_tag)

# ===============================
# 模块独立测试
# ===============================
"""
if __name__ == "__main__":
    log = Logger()
    log.info("系统初始化完成")
    log.warn("正在加载 Excel 模块 ...")
    log.error("文件未找到：data/巡检报告数据集.xlsx")
"""
# ========== End of 日志 Logger 类 ==========
