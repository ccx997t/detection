#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
巡检报告模板插入表格图片模块（report_embedder.py）
--------------------------------------------------
功能说明：
    本模块用于将 Excel 表格截图（JPG 文件）自动嵌入到 Word 巡检报告模板中，
    按照模板中定义的占位符（如 {{表1}}、{{表2}} ... {{表7}}）依次替换为对应的图片。

核心流程：
    1. 加载 Word 模板文件；
    2. 遍历模板中的段落与表格，查找占位符；
    3. 根据占位符名称加载对应目录下的 JPG 文件；
    4. 在占位符处插入图片（自动居中、宽度固定）；
    5. 保存生成的最终报告文件。

输入输出：
    - 输入：Word 模板文件路径、表格截图目录（IMAGES_DIR）
    - 输出：生成的完整巡检报告 Word 文件（保存在 OUTPUT_DIR）

依赖模块：
    - python-docx：Word 文档操作
    - util.Logger：自定义日志输出（来自 modules/util.py）
"""

import os
import re
import sys
from typing import Dict
from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate, InlineImage      # 导入 docxtpl 模板类与插图类
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import configparser
from jinja2 import Environment, DebugUndefined
# ============================================================
# 修正项目模块搜索路径，确保可导入 modules 下的工具模块
# ============================================================
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.append(PROJECT_ROOT)
print(f">> PROJECT_ROOT = {PROJECT_ROOT},  __file__ = {__file__}")

# ============================================================
# 项目模块 util
# ============================================================
try:
    from modules import util as _ut
except Exception as e:
    _ut = None
    print(f"⚠️  未找到 util 模块：{e}")

log = _ut.Logger()

# 全局参数
TEMPLATE_PATH = ""
PDFS_DIR = ""
IMAGES_DIR = ""
OUTPUT_DIR = ""
COVER_TEMPLATE_PATH = ""
# ============================================================
# 模块函数定义
# ============================================================

def load_template_file() -> Document:
    """加载 Word 模板文件，若失败则提示并退出。"""
    try:
        log.info(f"加载模板文件：{TEMPLATE_PATH}")
        doc = Document(TEMPLATE_PATH)
        log.info("✅ 模板加载成功。")
        return doc
    except Exception as e:
        log.error(f"❌ 模板文件加载失败：{TEMPLATE_PATH}，错误信息：{e}")
        sys.exit(1)


def load_jpg_files() -> Dict[str, str]:
    """
     加载指定目录下的所有表格截图文件，并构建占位符映射表。
     功能：
         - 自动扫描表1.jpg ~ 表7.jpg；
         - 生成占位符与对应图片路径的字典：
             例如：{"{{表1}}": "/path/to/表1.jpg", ...}
     参数：
         images_dir (str): 图片存放目录路径
     返回：
         Dict[str, str]: 占位符 → 图片路径 的映射表
     """
    log.info("✅ 加载图片文件.....")
    image_map = {}
    # 检查图片目录是否存在
    if not os.path.isdir(IMAGES_DIR):
        log.info(f"❌ 图片目录不存在：{IMAGES_DIR}")
        return  # 退出函数

    # 获取所有 .jpg 文件名列表
    image_files = [f for f in os.listdir(IMAGES_DIR) if f.lower().endswith(".jpg")]
    # 如果未发现任何图片，提示用户
    if not image_files:
        log.info("⚠️ 未找到任何 .jpg 文件，模板将不进行替换。")

    # 遍历所有图片文件
    for filename in image_files:
        key = os.path.splitext(filename)[0]  # 从文件名中提取变量名（去除扩展名）
        img_path = os.path.join(IMAGES_DIR, filename)  # 拼接图片的完整路径

        # 再次确认文件存在（保险处理）
        if os.path.exists(img_path):
            image_map[key] =img_path
            log.info(f"✅ 已准备图片：{img_path} → 模板变量 {{ {key} }}")  # 打印图片加载成功信息
    log.info("✅ 图片文件加载成功。")
    return image_map


def find_placeholders_and_replace_docxtemplate(doc: DocxTemplate, image_map: Dict[str, str]) -> None:
    """
    遍历整个文档（段落与表格单元格），匹配占位符并插入图片。
    逻辑：
        - 优先扫描所有段落；
        - 再扫描表格内的所有单元格；
        - 每当匹配到占位符（如 {{表3}}），则调用 replace_placeholder_with_image()。
    参数：
        DocxTemplate: Word 文档对象
        image_map (Dict[str, str]): 占位符 → 图片路径 映射表
    """

    context = {}  # 初始化上下文字典，用于存放变量名和图片对象
    for key, img_path in image_map.items():
        context[key] = InlineImage(doc, img_path, width=Inches(6.5))  # 设置图片宽度为 6.5 英寸
        log.info(f"键：{key}，值：{img_path}")
    # 创建 Jinja 环境对象
    jinja_env = Environment(undefined=DebugUndefined)

    # 渲染模板
    doc.render(context, jinja_env=jinja_env)


def find_placeholders_and_replace(doc: Document, image_map: Dict[str, str]) -> Document:
    """ 遍历整个文档（段落与表格单元格），匹配占位符并插入图片。
    逻辑：
        - 优先扫描所有段落；
        - 再扫描表格内的所有单元格；
        - 每当匹配到占位符（如 {{表3}}），则调用 replace_placeholder_with_image()。
    参数：
        doc (Document): Word 文档对象
        image_map (Dict[str, str]): 占位符 → 图片路径 映射表
    """
    # ---------- 1. 替换段落中的占位符 ----------
    for paragraph in doc.paragraphs:
        for placeholder, image_path in image_map.items():
            if placeholder in paragraph.text:
                log.info(f"匹配段落占位符：{placeholder}")
                #replace_placeholder_with_image(paragraph, image_path)

    # ---------- 2. 替换表格单元格中的占位符 ----------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, image_path in image_map.items():
                        if placeholder in paragraph.text:
                            log.info(f"匹配表格占位符：{placeholder}")
                            #replace_placeholder_with_image(paragraph, image_path)

    return doc


def clean_doc(doc: Document) -> Document:
    """ 清洗 doc对象，去除潜在的损坏段落或空元素。
    适用于 Word 打开时提示“内容有错误”的情况。
    """
    try:
        removed_count = 0

        # 清除段落中完全空的 run（无文本、无图片）
        for paragraph in doc.paragraphs:
            original_runs = paragraph.runs[:]
            for run in original_runs:
                if not run.text.strip() and not run._element.xpath(".//w:drawing"):
                    paragraph._element.remove(run._element)
                    removed_count += 1

        # 清除表格中空的段落
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_paragraphs = cell.paragraphs[:]
                    for p in original_paragraphs:
                        if not p.text.strip() and not p._element.xpath(".//w:drawing"):
                            cell._element.remove(p._element)
                            removed_count += 1

        log.info(f"✅ 清洗完成，移除空 run/段落共 {removed_count} 个元素。")
    except Exception as e:
        log.error(f"❌ 清洗 doc 对象失败：{e}")
    return doc

def create_report_cover(config : configparser.ConfigParser(), info: dict):
    """
    生成巡检报告封面。
    输出路径：out/实验性项目巡检报告.docx
    """
    TEMPLATE_PATH =config.get("Path", "template_path")
    basename = os.path.basename(TEMPLATE_PATH)
    #print(f"basename = {basename}")
    # 去掉文件名中的“模板”，构成输出文件名。
    new_name = re.sub(r"模板\(.*?\)", "", basename).replace(".docx", "")
    #print(f"new_name = {new_name}")
    new_name = new_name.strip("-_ ") + ".docx"
    #print(f"new_name = {new_name}")
    # 构成输出文件全路径。
    OUTPUT_DIR = config.get("Path", "output_dir")
    log.info(f"📄 OUTPUT_DIR：{OUTPUT_DIR}")
    output_path = os.path.join(OUTPUT_DIR, new_name)
    log.info(f"📄 正在生成封面：{output_path}")
    # 填充模板上下文
    context = {
        "项目名称": info.get("project_name", ""),
        "机房名称": info.get("room_name", ""),
        "年度": info.get("year", ""),
        "季度": info.get("quarter", ""),
        "报告日期": info.get("report_date",),
        "责任人": info.get("report_person", ""),
    }
    jinja_env = Environment(undefined=DebugUndefined)
    doc = DocxTemplate(output_path)
    doc.render(context, jinja_env=jinja_env)
    doc.save(output_path)
    log.info(f"✅ 封面生成成功：{output_path}")

def save_doc(doc: DocxTemplate) -> str:
    """ 保存 Word 文档到指定目录。 """
    # 从模板弯路路径取出模板文件名
    basename = os.path.basename(TEMPLATE_PATH)
    #print(f"basename = {basename}")
    # 去掉文件名中的“模板”，构成输出文件名。
    new_name = re.sub(r"模板\(.*?\)", "", basename).replace(".docx", "")
    #print(f"new_name = {new_name}")
    new_name = new_name.strip("-_ ") + ".docx"
    #print(f"new_name = {new_name}")
    # 构成输出文件全路径。
    output_path = os.path.join(OUTPUT_DIR, new_name)
    # 确保输出文件目录存在。
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    # 保存输出文件。
    doc.save(output_path)
    log.info(f"✅ 生成报告成功：{output_path}")
    doc = Document(output_path)
    doc.save(output_path)
    """
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    doc = word.Documents.Open(output_path, ConfirmConversions=False, ReadOnly=False)
    repaired = output_path.replace(".docx", "_fixed.docx")
    doc.SaveAs(repaired, FileFormat=16)  # 16 = wdFormatXMLDocument
    doc.Close()
    word.Quit()
    """
    return output_path


def run(config: configparser.ConfigParser):
    """ 模块主执行函数。 """
    # 提取配置文件参数项
    global TEMPLATE_PATH, IMAGES_DIR, OUTPUT_DIR
    TEMPLATE_PATH = config.get("Path", "template_path")
    IMAGES_DIR = config.get("Path", "images_dir")
    OUTPUT_DIR = config.get("Path", "output_dir")
    # 加载模板。
    doc = DocxTemplate(TEMPLATE_PATH)  # 加载 Word 模板为 docxtpl 文档对象
    # 加载表格截图映射表；
    image_map = load_jpg_files()
    # 查找占位符并替换为图片
    find_placeholders_and_replace_docxtemplate(doc, image_map)
    # 保存生成的新报告文件
    save_doc(doc)

# ============================================================
# 测试运行（仅在独立运行时触发）
# ============================================================
if __name__ == "__main__":
    TEMPLATE_PATH = "../template/实验性项目巡检报告模板(1.0).docx"
    IMAGES_DIR = "../tmp/images/"
    OUTPUT_DIR = "../out/"
    # 加载模板。
    doc = DocxTemplate(TEMPLATE_PATH)  # 加载 Word 模板为 docxtpl 文档对象
    #加载表格截图映射表；
    image_map = load_jpg_files()
    #查找占位符并替换为图片
    find_placeholders_and_replace_docxtemplate(doc, image_map)
    #保存生成的新报告文件
    save_doc(doc)
    # info = {
    #         "project_name": "实验性项目AI巡检系统",
    #         "room_name": "主数据中心机房",
    #         "year": 2025,
    #         "quarter": "Q4",
    #         "report_date": "2025年3月",
    #         "report_person": "张三"
    # }
    # create_report_cover(999,info)
